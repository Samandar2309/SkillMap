from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

A = "ʻ"  # oʻ, gʻ
E = "ʼ"  # taʼlim, maʼlumot, sunʼiy

BASE_DIR = Path(__file__).resolve().parent
SOURCE_CANDIDATES = [
    BASE_DIR / "DL Jumabayev Samandar.docx",
    BASE_DIR / "DL_Jumabayev_Samandar_v2.docx",
]
OUTPUT = BASE_DIR / "DL_Jumabayev_Samandar_120bet_final.docx"


def _set_font(run, *, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), "Times New Roman")


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_para(doc, text, *, bold=False, italic=False, align="justify", first_line=True, size=14, after=6, before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(after)
    if before:
        pf.space_before = Pt(before)
    if first_line:
        pf.first_line_indent = Cm(1.25)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, *, level=1, align="center"):
    size = 16 if level == 1 else 14
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    return p


def add_caption(doc, text):
    return add_para(doc, text, align="center", italic=True, first_line=False, size=12, after=4)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        _set_font(run, size=12, bold=True)
        _set_cell_border(cell)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            _set_font(run, size=12)
            _set_cell_border(cell)
    return table


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.left_indent = Cm(1.0)
        run = p.add_run(line if line else " ")
        _set_font(run, size=11)
        run.font.name = "Consolas"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "cs"):
            rFonts.set(qn(f"w:{attr}"), "Consolas")


def pick_source() -> Path:
    for path in SOURCE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("Asosiy DOCX topilmadi")


CH4_INTRO = [
    f"Mazkur bob diplom loyihasining amaliy-tashkiliy qismiga bagʻishlanadi va SkillMap platformasini ishlab chiqish hamda undan foydalanish jarayonida inson salomatligi, mehnat muhofazasi, texnik xavfsizlik va axborot xavfsizligiga oid masalalarni tizimli ko'rib chiqadi. Dasturiy mahsulot faqat funksional jihatdan emas, balki foydalanuvchi uchun xavfsiz, ergonomik va uzoq muddatli foydalanishga yaroqli bo'lishi bilan ham baholanadi. Shu sababli bobda kompyuter bilan ishlash sharoitlari, server va tarmoq infratuzilmasining xavfsizligi, favqulodda holatlarga tayyorgarlik hamda AI xizmatlaridan mas'uliyatli foydalanish tamoyillari alohida tahlil qilinadi.",
    f"SkillMap kabi zamonaviy platformalar ko'pincha bir nechta muhitda ishlaydi: foydalanuvchi kompyuteri, brauzer, mobil qurilma, backend server, ma'lumotlar bazasi va tashqi AI provayder. Har bir qatlam o'ziga xos xavfga ega. Foydalanuvchi tomonida ko'z charchashi, noto'g'ri o'tirish holati yoki kiberhujumlar xavfi mavjud bo'lsa, server tomonida elektr uzilishi, ma'lumotlar yo'qolishi, noto'g'ri konfiguratsiya yoki ruxsatsiz kirish tahdidi paydo bo'ladi. Demak, hayot faoliyati xavfsizligi masalasi faqat ishlab chiqarish binolariga emas, balki raqamli ta'lim tizimiga ham taalluqlidir.",
]

CH4_1_PARAS = [
    f"4.1. Kompyuter xavfsizligi va sanitariya-gigiyena qoidalari",
    f"Ta'lim va dasturiy ishlab chiqish bilan bog'liq ish joyi bir vaqtning o'zida ham qulay, ham xavfsiz bo'lishi kerak. Foydalanuvchi, dasturchi yoki administrator kun davomida ekran qarshisida uzoq o'tirganda, tananing tabiiy holati buziladi, ko'z mushaklari zo'riqadi va diqqat pasayadi. Shuning uchun ish joyini tashkil etishda monitor balandligi, klaviatura joylashuvi, o'rindiqning tayanchi, yoritish darajasi va xonadagi havo almashinuvi alohida e'tiborga olinadi. Monitor ko'zdan taxminan 50–70 sm masofada joylashishi, ekran yuqori qismi ko'z darajasidan biroz pastroqda bo'lishi va yorug'lik manbalari bevosita ekranga aks etmasligi tavsiya etiladi.",
    f"Sanitariya-gigiyena nuqtai nazaridan kompyuter bilan ishlashdagi eng muhim talablaridan biri — uzluksiz ish vaqtini to'g'ri taqsimlashdir. Uzoq vaqt tanaffussiz ishlash bosh og'rig'i, ko'z qurishi va bo'yin mushaklarida kuchlanish keltirib chiqaradi. Shu sababli har 45–60 daqiqada qisqa tanaffus qilish, ko'z uchun mashqlar bajarish va tanani cho'zish foydali hisoblanadi. Amaliyotda 20-20-20 qoidasi yaxshi natija beradi: har 20 daqiqada 20 soniya davomida kamida 20 fut masofadagi ob'ektga qarash ko'z zo'riqishini kamaytiradi. SkillMap foydalanuvchiga kunlik vazifalarni rejalashtirar ekan, aynan mana shunday mikro-tanaffuslarni ham e'tiborga oladi.",
    f"Ergonomika mehnat unumdorligi bilan bevosita bog'liq. Noto'g'ri stul yoki klaviatura joylashuvi faqat noqulaylik tug'dirmaydi, balki uzoq muddatda tayanch-harakat apparati kasalliklariga sabab bo'lishi mumkin. Shu bois ish o'rnida tirsak burchagi taxminan 90 daraja bo'lishi, bilaklar neytral holatda saqlanishi va oyoqlar polga to'liq tegib turishi kerak. Agar foydalanuvchi noutbukdan foydalansa, alohida klaviatura va stend qo'llash maqsadga muvofiqdir. Bu oddiy tavsiyalar amalda ishga diqqatni oshirib, charchoqni kamaytiradi.",
    f"Kompyuter xavfsizligi faqat jismoniy qulaylikdan iborat emas. Raqamli muhitda parollar, brauzer sozlamalari, antivirus dasturlari, brauzer kengaytmalari va fayllar bilan ehtiyotkor ishlash ham zarur. Foydalanuvchi bir xil parolni barcha xizmatlarda qo'llamasligi, kuchli parol siyosatiga rioya qilishi va imkon bo'lsa ikki bosqichli autentifikatsiyani yoqishi kerak. SkillMap platformasida bu yondashuv JWT tokenlar, qisqa muddatli access tokenlar va rollarga asoslangan kirish nazorati orqali qo'llab-quvvatlanadi.",
    f"Kiberxavfsizlik bo'yicha eng ko'p uchraydigan tahdidlardan biri phishing hisoblanadi. Noma'lum manbadan kelgan havolalar, soxta login sahifalar va shubhali ilovalar foydalanuvchi ma'lumotlarini o'g'irlashga qaratilgan bo'lishi mumkin. Shu sababli foydalanuvchini o'qitish ham xavfsizlikning bir qismidir. Platforma ichida foydalanuvchiga parolni hech kimga bermaslik, umumiy kompyuterlarda sessiyani yopish va shubhali xatlar bo'lsa administratorga xabar berish haqida eslatmalar chiqariladi.",
    f"Axborot xavfsizligi nuqtai nazaridan ma'lumotlarning tranzit va saqlash vaqtida himoyalanishi birlamchi shartdir. SkillMap serverlari HTTPS protokoli orqali ishlaydi, API so'rovlari CSRF va CORS siyosatlari bilan cheklanadi, ma'lumotlar bazasiga kirish esa qat'iy ruxsatlar bilan boshqariladi. Ma'lumotlarning zaxira nusxalari muntazam yaratilishi, ularni tiklash testi o'tkazilishi va log fayllari tahlil qilinishi talaba ma'lumotlarining yo'qolish xavfini kamaytiradi.",
    f"Raqamli gigiyena ham muhim. Brauzerda keraksiz qo'shimcha kengaytmalar o'rnatish, bir vaqtning o'zida juda ko'p tab ochish yoki notanish qurilmadan tizimga kirish xavf tug'diradi. Shu sababli platforma foydalanuvchini faqat ishonchli qurilmalardan foydalanishga, jamoat kompyuterlarida sessiyani to'g'ri yakunlashga va avtomatik to'ldirish funksiyalarini ehtiyotkorlik bilan qo'llashga undaydi. Bularning barchasi oddiy ko'rinsa-da, amalda ma'lumotlar xavfsizligining birinchi qatlamini tashkil etadi.",
    f"SkillMap tizimida foydalanuvchining sog'lig'iga oid masalalar ham e'tiborga olingan. Interfeysdagi ranglar kontrasti, matn hajmi, tugmalar orasidagi masofa va xabarlarning o'qilishi osonligi foydalanuvchi charchoqlarini kamaytiradi. Sahifalar ko'p vizual shovqindan xoli bo'lishi kerak: juda yorqin gradientlar yoki keraksiz animatsiyalar uzoq ishlashda chalg'itadi. Shuning uchun dizayn tamoyili minimalizm, aniq ierarxiya va tez o'qiladigan matnga tayanadi.",
]

CH4_1_TABLE = [
    ["Ko'rinadigan xavf", "Ehtimoliy oqibat", "Profilaktik chora"],
    ["Noto'g'ri o'tirish holati", "Bo'yin va bel og'rig'i", "Ergonomik o'rindiq, monitor balandligini moslash"],
    ["Uzluksiz uzoq ish", "Ko'z charchashi va diqqat pasayishi", "Har 45–60 daqiqada tanaffus va ko'z mashqlari"],
    ["Phishing havolalari", "Hisob buzilishi", "Ikki bosqichli autentifikatsiya va xabarlarni tekshirish"],
    ["Zaif parollar", "Ruxsatsiz kirish", "Murakkab parol siyosati va parol menejeri"],
    ["Noto'g'ri brauzer kengaytmasi", "Ma'lumotlarning sizib chiqishi", "Faqat ishonchli kengaytmalarni o'rnatish"],
    ["Shifrlanmagan trafik", "So'rovlarni ushlab olish", "HTTPS va xavfsiz cookie sozlamalari"],
    ["Zaxira nusxasining yo'qligi", "Ma'lumotni tiklab bo'lmaslik", "Muntazam backup va tiklash testi"],
]

CH4_2_PARAS = [
    f"4.2. Texnik obyektlarda mehnatni muhofaza qilishni tashkil etish",
    f"SkillMap platformasi bevosita ishlab chiqarish sexi yoki qurilish maydonida yaratilmasa ham, uning server xonasi, ishchi stansiyalar va tarmoq uskunalari texnik obyekt sifatida qaraladi. Bunday obyektlarda mehnatni muhofaza qilishning maqsadi — xodimlar, talaba-praktikantlar va texnik infratuzilmaning xavfsiz ishlashini ta'minlashdir. Server xonasida havo harorati, namlik, chang darajasi, elektr ta'minoti, yong'in xavfsizligi va kirish nazorati alohida boshqarilishi kerak.",
    f"Elektr xavfsizligi birinchi o'rinda turadi. Kompyuterlar, marshrutizatorlar, UPS qurilmalari va quvvat manbalari bilan ishlaganda yerga ulash, kuchlanishdan himoya va ortiqcha yuklamani cheklash zarur. Kabel izolyatsiyasi shikastlangan bo'lsa, u darhol almashtiriladi. Tarmoq shkafida tartibsiz kabellar faqat estetik muammo emas, balki qisqa tutashuv va xizmat ko'rsatishdagi xatolar manbaidir. Shu sababli kabel markirovkasi, ranglar bo'yicha ajratish va kabel-kanallardan foydalanish tavsiya etiladi.",
    f"Elektr uzilishi yuz berganda ma'lumotlarning yo'qolish ehtimoli yuqori bo'ladi. Shu sababli server va routerlar uchun uzluksiz elektr ta'minoti manbasi (UPS) qo'llanadi, muhim xizmatlar esa avtomatik ravishda zaxira quvvatga o'tadi. Faqat quvvatning saqlanishi yetarli emas — tizim xavfsiz o'chishi, tranzaksiyalar to'g'ri yakunlanishi va baza yaxlitligi buzilmasligi ham kerak. Shu maqsadda zaxira nusxalash rejasi va replikatsiya mexanizmlari ishlab chiqiladi.",
    f"Yong'in xavfsizligi nuqtai nazaridan server xonasi doimo toza, quruq va nazorat ostida bo'lishi kerak. Elektr jihozlari yonida yonuvchan materiallar saqlanmasligi, evakuatsiya yo'llari ochiq bo'lishi va o'chirish vositalari — ko'pikli yoki karbonat angidridli o't o'chirgichlar — tayyor turishi kerak. Xodimlar favqulodda vaziyatda qaysi tugmani bosish, qaysi kontaktni uzish va qaysi xavfsiz yo'nalish bo'ylab chiqish kerakligini oldindan bilishi lozim.",
    f"Mehnat muhofazasini tashkil etishda hujjatlashtirish muhim o'rin tutadi. Har bir texnik xodim uchun ish yo'riqnomasi, xavfsizlik bo'yicha dastlabki va davriy o'qitish, jurnal qaydlari va hodisalar ro'yxati bo'lishi lozim. SkillMap kabi platformada bu jarayon administratorlar uchun alohida rol va ruxsatlar orqali boshqariladi. Tizim sozlamalaridagi o'zgarishlar audit logda saqlanadi, muhim amallar esa vaqt tamg'asi bilan qayd etiladi.",
    f"Texnik xizmat ko'rsatish jarayonida ham ehtiyot choralariga rioya qilish kerak. Serverga bog'langan tarmoq kabelini uzishdan oldin foydalanuvchi sessiyalari yakunlanganini tekshirish, ma'lumotlar bazasini backup qilish va xizmatlar o'chirilganini tasdiqlash zarur. Buyruq satri orqali noto'g'ri operatsiya bajarilishi yoki muhim fayllar tasodifan o'chirilishi xavfi mavjud. Shu bois ishlab chiqish va production muhitlari aniq ajratiladi, test ma'lumotlari real ma'lumotlar bilan aralashtirilmaydi.",
    f"Favqulodda vaziyatlar rejasi har bir tashkilotda bo'lishi kerak. Bu reja elektr ta'minoti uzilishi, tarmoqdagi nosozlik, ma'lumotlar bazasi buzilishi, AI provayderining javob bermasligi yoki kiberhujum yuz berganda qanday harakat qilish kerakligini belgilaydi. SkillMap uchun birinchi darajadagi choralar — xizmatni qayta ishga tushirish, loglarni tahlil qilish, zaxira provayderga o'tish va foydalanuvchi ma'lumotlarini tiklashdan iborat. Ikkinchi darajali choralar esa muammo sababini bartaraf etish va kelajakda takrorlanmasligi uchun konfiguratsiyani tuzatishdir.",
    f"Axborot xavfsizligini mehnat muhofazasining ajralmas qismi sifatida ko'rish kerak. Chunki foydalanuvchi ma'lumotlari yo'qolsa yoki buzilsa, bu nafaqat texnik, balki pedagogik va huquqiy zarar ham keltiradi. Shu sababli ruxsat darajalari, shifrlash, ma'lumotlarni saqlash muddati, foydalanuvchi roziligi va ma'lumotlarni o'chirish huquqi aniq belgilangan. Platformada administrator faqat zarur bo'lgan ma'lumotlarni ko'radi, tizim esa minimal ma'lumot yig'ish prinsipi asosida ishlaydi.",
    f"Texnik obyektlarda sog'liqni saqlash bilan bog'liq qo'shimcha omillar ham mavjud. Xona shovqini, havoning quruqligi, konditsionerning noto'g'ri ishlashi yoki noto'g'ri yoritish xodimning ish qobiliyatiga ta'sir qiladi. Shu sababli server xonasi va ish joylarida sanitariya nazorati muntazam o'tkaziladi. O'sha muhitdagi harorat va namlik me'yoriy diapazonda ushlab turilishi, chang filtrlar vaqtida tozalanishi va favqulodda holatdagi aloqa telefonlari ko'rinadigan joyda bo'lishi kerak.",
]

CH4_2_TABLE = [
    ["Tashkiliy yo'nalish", "Amaliy chora", "Kutiladigan natija"],
    ["Elektr xavfsizligi", "Yerga ulash va UPS", "Kuchlanishdan himoya"],
    ["Yong'in xavfsizligi", "O't o'chirgich va evakuatsiya rejasi", "Favqulodda chiqish tezlashadi"],
    ["Kabel boshqaruvi", "Markirovka va kabel-kanal", "Qisqa tutashuv xavfi kamayadi"],
    ["Zaxira nusxa", "Avtomatik backup va tiklash testi", "Ma'lumot yo'qolishi kamayadi"],
    ["Kirish nazorati", "RBAC va audit log", "Ruxsatsiz kirish cheklanadi"],
    ["Xodimlarni o'qitish", "Yo'riqnoma va davriy trening", "Xavfsizlik madaniyati oshadi"],
    ["Incident response", "Ssenariy va aloqa zanjiri", "Nosozlikka javob tezlashadi"],
]

CH4_3_PARAS = [
    f"Alohida ahamiyatga ega masala — AI xizmatlaridan xavfsiz foydalanishdir. SkillMap platformasi foydalanuvchi rejasini generatsiya qilishda tashqi katta til modellari bilan ishlaydi. Bu esa ikki turdagi xavfni keltirib chiqaradi: texnik xavf va mazmuniy xavf. Texnik xavf API kvotalari tugashi, kechikish, tarmoq uzilishi yoki noto'g'ri autentifikatsiya ko'rinishida namoyon bo'ladi. Mazmuniy xavf esa modelning xato tavsiya berishi, mavjud bo'lmagan resursni to'qib chiqarishi yoki noto'g'ri ketma-ketlik taklif qilishida ko'rinadi. Shu sababli tizimda validator va anti-hallucination qatlami joriy etilgan.",
    f"AI chiqishini avtomatik tekshirish foydalanuvchini noto'g'ri yo'nalishlardan himoya qiladi. Agar model tomonidan tavsiya etilgan mavzu oldingi kompetensiyalar zanjiriga mos kelmasa, reja qayta tuziladi yoki zaxira prompt ishga tushiriladi. Bu yondashuv ayniqsa ta'lim tizimida muhim: talaba vaqtini noto'g'ri kontentga sarflamasligi kerak. Shuning uchun AI yordamchi rolini bajaradi, yakuniy qaror esa platforma mantiqi va foydalanuvchi nazorati ostida qoladi.",
    f"Prompt injection va nojo'ya so'rovlar ham zamonaviy AI xavfsizligining bir qismidir. Foydalanuvchi modelga tizim qoidalarini buzishga urinadigan matn yuborishi mumkin. Bunga qarshi qat'iy prompt shablonlari, konteksni tozalash va javobni sxema bo'yicha cheklash usullari qo'llaniladi. SkillMap arxitekturasida AI-ga yuboriladigan so'rovlar oldindan formatlanadi, foydalanuvchi kiritgan erkin matn esa minimal darajada ishlatiladi.",
    f"Ma'lumotlarni maxfiy saqlash ham AI bilan bog'liq muhim talabdir. Talabaning maqsadi, qiziqishlari va akademik natijalari maxfiy hisoblanadi. Tizim bu ma'lumotlarni faqat zarur miqdorda yig'adi, ortiqcha maydonlar so'ramaydi va ularni ruxsat darajasi orqali himoya qiladi. Administrator va tahlilchi rollari bir xil huquqqa ega emas; bu farq RBAC modelida aniq belgilanadi. Natijada shaxsiy ma'lumotlarga kirish nazorat ostida bo'ladi.",
    f"Mavjudlik va ishonchlilik nuqtai nazaridan bir necha qo'shimcha choralar zarur. AI servisi vaqtincha ishlamasa, foydalanuvchi butun platformadan foydalanishni to'xtatmasligi kerak. Shu sababli fallback generator mavjud bo'lib, u oldindan yaratilgan shablonlar asosida soddalashtirilgan yo'l xaritasini qaytaradi. Bu yo'l xaritasi ideal bo'lmasligi mumkin, biroq foydalanuvchini bloklab qo'ymaydi va tizimning uzluksiz ishlashini ta'minlaydi.",
    f"Monitoring va kuzatuv platforma xavfsizligining so'nggi qatlamidir. Xatoliklar loglarda saqlanadi, AI javoblarining sifati test jarayonida nazorat qilinadi va shubhali so'rovlar alohida qayd etiladi. Agar bir xil foydalanuvchidan qisqa vaqt ichida juda ko'p talab yuborilsa, rate limiting qo'llanadi. Bu choralar tizimni suiiste'moldan himoya qiladi va xizmat sifatini barqaror saqlaydi.",
]

CH4_SUMMARY = [
    f"Hayot faoliyati xavfsizligi bo'yicha xulosa shundan iboratki, SkillMap platformasi faqat dasturiy jihatdan emas, balki inson salomatligi va axborot xavfsizligi nuqtai nazaridan ham puxta rejalashtirilishi kerak. Ish joyining ergonomik tashkil etilishi, server xonasining elektr va yong'in xavfsizligi, ma'lumotlarni zaxiralash, AI javoblarini tekshirish hamda foydalanuvchi ma'lumotlarini himoya qilish bir-birini to'ldiruvchi choralar hisoblanadi. Ushbu talablar bajarilganda platformadan foydalanish qulay, xavfsiz va uzoq muddatli bo'ladi.",
]

XULOSA_PARAS = [
    f"Mazkur diplom loyihasida sun'iy intellekt yordamida talabalar shaxsiy rivojlanish rejasini avtomatik tuzuvchi SkillMap platformasini ishlab chiqish masalasi kompleks tarzda ko'rib chiqildi. Ishning kirish qismida muammoning dolzarbligi, maqsadi, vazifalari va amaliy ahamiyati asoslandi. Birinchi bobda shaxsiy rivojlanish tushunchasi, ta'limdagi sun'iy intellekt texnologiyalari va mavjud platformalar qiyosiy tahlil qilinib, yangi yechim uchun nazariy asos yaratildi.",
    f"Ikkinchi bobda foydalanuvchi profilini yig'ish, kompetensiya bo'shliqlarini aniqlash, tavsiyalarni saralash va rivojlanish rejasini generatsiya qilish algoritmlari ishlab chiqildi. Ushbu bosqichda kompetensiya profili, yaqin rivojlanish zonasi, gibrid filtrlash va izohlanuvchan sun'iy intellekt tamoyillari bir tizimda birlashtirildi. Natijada talabaning maqsadi, vaqt resursi va o'rganish uslubiga mos yo'l xaritasini shakllantirish mexanizmi yaratildi.",
    f"Uchinchi bobda platformaning real dasturiy amalga oshirilishi yoritildi. Django, DRF, Celery, React, TypeScript, TailwindCSS va tashqi LLM provayderlaridan foydalanilgan holda ishlovchi prototipning arxitekturasi, modul tuzilishi va API endpointlari tavsiflandi. Sinov natijalari tizimning funksional to'g'riligi, javob tezligi va foydalanuvchi uchun qulayligi amaliy jihatdan qoniqarli ekanini ko'rsatdi. Ayniqsa AI qatlamida anti-hallucination va fallback mexanizmlarining mavjudligi tizimning ishonchliligini oshirdi.",
    f"To'rtinchi bobda hayot faoliyati xavfsizligi, sanitariya-gigiyena, elektr xavfsizligi, yong'in xavfsizligi, mehnat muhofazasi va axborot xavfsizligi masalalari yoritildi. Bu qismda nafaqat server xonasi va texnik infratuzilma uchun, balki kundalik kompyuter bilan ishlash sharoitlari uchun ham aniq tavsiyalar berildi. Shu bilan loyiha inson salomatligi, ma'lumotlar maxfiyligi va tizim barqarorligi nuqtai nazaridan to'liqroq ko'rinish oldi.",
    f"Umumiy xulosa sifatida aytish mumkinki, SkillMap platformasi talabalar uchun shaxsiy rivojlanish rejasini avtomatik shakllantirish, uni dinamik yangilash va foydalanuvchiga tushunarli izoh berish imkonini yaratadi. Ushbu yechim o'zbek tilida ishlovchi, mahalliy kontekstga mos, kengaytirishga tayyor va amaliy qo'llashga yaqin bo'lgan raqamli platforma sifatida baholandi. Kelgusida tizimni universitetlarning rasmiy axborot tizimlari, mobil ilova va ish bozori ma'lumotlari bilan chuqur integratsiya qilish maqsadga muvofiqdir.",
    f"Shaxsiy tavsiya sifatida, platformani keyingi bosqichda ko'proq real foydalanuvchilar guruhi bilan sinovdan o'tkazish, o'zbek tilidagi kontent bazasini kengaytirish, foydalanuvchi fikriga asoslangan adaptiv mexanizmlarni kuchaytirish va o'qituvchi panelini boyitish taklif etiladi. Ushbu yo'nalishlar amalga oshirilsa, SkillMap nafaqat diplom loyihasi, balki amaliy raqamli mahsulot sifatida ham rivojlanishi mumkin.",
]

REFERENCES = [
    "1. Oʻzbekiston Respublikasi Konstitutsiyasi. — T., 2023.",
    "2. Oʻzbekiston Respublikasining «Taʼlim toʻgʻrisida»gi Qonuni, OʻRQ-637. — 2020.",
    "3. Oʻzbekiston Respublikasining «Shaxsiy maʼlumotlar toʻgʻrisida»gi Qonuni, OʻRQ-547. — 2019.",
    "4. Oʻzbekiston Respublikasining «Axborotlashtirish toʻgʻrisida»gi Qonuni, 560-II. — 2003.",
    "5. Oʻzbekiston Respublikasi Prezidentining PF-6079-son Farmoni «Raqamli Oʻzbekiston — 2030» strategiyasi toʻgʻrisida. — 2020.",
    "6. Oʻzbekiston Respublikasi Prezidentining PQ-358-son Qarori «Sunʼiy intellekt texnologiyalarini joriy etish strategiyasi» toʻgʻrisida. — 2024.",
    "7. Oʻzbekiston Respublikasi Prezidentining oliy taʼlimni rivojlantirishga oid farmon va qarorlari. — T., 2017–2019.",
    "8. ISO/IEC 27001:2022. Information Security Management Systems — Requirements.",
    "9. ISO/IEC 25010:2011. Systems and Software Engineering — Systems and Software Quality Requirements and Evaluation.",
    "10. UNESCO. Recommendation on the Ethics of Artificial Intelligence. — 2021.",
    "11. Corbett A. T., Anderson J. R. Knowledge Tracing: Modeling the Acquisition of Procedural Knowledge. — 1994.",
    "12. Piech C. et al. Deep Knowledge Tracing. — 2015.",
    "13. Vygotsky L. S. Mind in Society: The Development of Higher Psychological Processes. — 1978.",
    "14. Fleming N. VARK: A Guide to Learning Styles. — 2001.",
    "15. Kolb D. A. Experiential Learning: Experience as the Source of Learning and Development. — 1984.",
    "16. Zimmerman B. J. Becoming a Self-Regulated Learner: An Overview. — 2002.",
    "17. Wiley J. & Sons. Adaptive Learning and Educational Technology Reference Materials. — 2023.",
    "18. Coursera Inc. Coursera for Campus product documentation. — 2024.",
    "19. LinkedIn Learning documentation and course recommendation overview. — 2024.",
    "20. EAB Navigate student success platform documentation. — 2024.",
    "21. Knewton adaptive learning platform technical overview. — 2024.",
    "22. Realizeit AI Graph methodology overview. — 2024.",
    "23. Django Software Foundation. Django Documentation. — 2025.",
    "24. Django REST Framework. Official Documentation. — 2025.",
    "25. React Team. React Documentation. — 2025.",
    "26. TypeScript Team. TypeScript Handbook. — 2025.",
    "27. Tailwind Labs. Tailwind CSS Documentation. — 2025.",
    "28. NIST. Secure Software Development Framework (SSDF). — 2022.",
    "29. OWASP Foundation. OWASP Top 10 Web Application Security Risks. — 2021.",
    "30. Microsoft. UX and Accessibility Design Guidelines. — 2024.",
]

APPENDIX_A = [
    f"Ilova A. SkillMap platformasini ishga tushirish va sozlash bo'yicha qisqa yo'riqnoma",
    f"Ushbu ilovada platformani lokal muhitda ishga tushirish uchun bajariladigan amallar ketma-ket bayon qilinadi. Amaliyotda ishlab chiquvchi birinchi navbatda Python virtual muhitini faollashtiradi, so'ng Django migratsiyalarini bajaradi va frontend uchun NPM paketlarini o'rnatadi. Barcha bosqichlar qayta takrorlanishi oson bo'lishi uchun ilovada qisqa, ammo aniq texnik ko'rsatmalar jamlandi.",
    f"1-qadam. Backend muhitini tayyorlash: `python -m venv .venv`, so'ng `pip install -r requirements.txt`. 2-qadam. Ma'lumotlar bazasi migratsiyasi: `python manage.py migrate`. 3-qadam. Administrator akkauntini yaratish: `python manage.py createsuperuser`. 4-qadam. Celery worker va beat jarayonlarini ishga tushirish. 5-qadam. Frontend tomonida `npm install` va `npm run dev` orqali ishga tushirish. 6-qadam. `.env` fayllardagi API manzillarini bir-biriga moslashtirish.",
]

APPENDIX_B = [
    f"Ilova B. Asosiy API endpointlari va ularning vazifalari",
]

APPENDIX_C = [
    f"Ilova C. Onboarding savollari va baholash mezonlari",
]

APPENDIX_D = [
    f"Ilova D. Xavf va sinovlar uchun tekshiruv ro'yxati",
]

API_ROWS = [
    ["POST", "/api/v1/register", "Yangi foydalanuvchi yaratish", "Yo'q"],
    ["POST", "/api/v1/login", "JWT access va refresh token olish", "Yo'q"],
    ["POST", "/api/v1/refresh", "Access tokenni yangilash", "Yo'q"],
    ["GET", "/api/v1/me", "Joriy foydalanuvchi ma'lumotini olish", "Ha"],
    ["POST", "/api/v1/onboarding/step-1", "Yo'nalish va maqsadni kiritish", "Ha"],
    ["POST", "/api/v1/onboarding/step-2", "Boshlang'ich bilim va tajriba", "Ha"],
    ["POST", "/api/v1/onboarding/step-3", "Vaqt byudjeti va ritm", "Ha"],
    ["POST", "/api/v1/onboarding/step-4", "O'rganish uslubi va til afzalligi", "Ha"],
    ["GET", "/api/v1/roadmap", "Saqlangan yo'l xaritasi", "Ha"],
    ["POST", "/api/v1/roadmap/generate", "AI orqali yangi yo'l xaritasi yaratish", "Ha"],
    ["GET", "/api/v1/tasks/today", "Bugungi vazifalar", "Ha"],
    ["POST", "/api/v1/tasks/complete", "Vazifani bajarilgan deb belgilash", "Ha"],
    ["GET", "/api/v1/progress", "Progress va streak ko'rsatkichlari", "Ha"],
    ["GET", "/api/v1/analytics/dashboard", "Analitika paneli", "Ha"],
    ["POST", "/api/v1/ai/analyze", "Profilni AI tahlil qilish", "Ha"],
    ["POST", "/api/v1/ai/motivate", "Motivatsion xabar generatsiyasi", "Ha"],
    ["GET", "/api/v1/docs/", "Swagger UI hujjati", "Yo'q"],
]

ONBOARDING_ROWS = [
    ["1", "Joriy yo'nalish, kurs va fakultet", "Bazaviy kontekst", "1–5"],
    ["2", "Asosiy maqsad va kutilgan natija", "SMART maqsadlash", "1–5"],
    ["3", "Bilim darajasi va tajriba", "Boshlang'ich profil", "0.0–1.0"],
    ["4", "Haftalik vaqt byudjeti", "Reja uzunligi", "15–240 daqiqa"],
    ["5", "O'rganish uslubi", "VARK moslashuvi", "Vizual/Audial/…"],
    ["6", "Afzal til va kontent turi", "Mahalliylashtirish", "UZ/RU/EN"],
    ["7", "Qo'shimcha qiziqishlar", "Resurs tanlash", "Matnli"],
    ["8", "Cheklovlar va noqulayliklar", "Yukni kamaytirish", "Erkin matn"],
    ["9", "Kasbiy yo'nalish", "Mehnat bozori mosligi", "Erkin matn"],
    ["10", "Motivatsion omillar", "Reja barqarorligi", "1–5"],
]

TEST_ROWS = [
    ["1", "Login/registratsiya ssenariysi", "JWT tokenlar to'g'ri beriladi", "O'tdi"],
    ["2", "Onboarding bosqichlari", "Ma'lumotlar yo'qolmay saqlanadi", "O'tdi"],
    ["3", "Roadmap generatsiyasi", "AI javobi validatsiyadan o'tadi", "O'tdi"],
    ["4", "Fallback rejimi", "Provayder ishlamasa zaxira yoqiladi", "O'tdi"],
    ["5", "Vazifa belgilash", "Progress yangilanadi", "O'tdi"],
    ["6", "Quiz yuborish", "Javoblar to'g'ri baholanadi", "O'tdi"],
    ["7", "Streak hisoblash", "Ketma-ketlik aniqlanadi", "O'tdi"],
    ["8", "Analitika paneli", "Grafiklar xato chiqarmaydi", "O'tdi"],
    ["9", "Ruxsat nazorati", "RBAC cheklovlari ishlaydi", "O'tdi"],
    ["10", "Audit log", "Muhim amallar qayd etiladi", "O'tdi"],
]

RISK_ROWS = [
    ["API kvotasi tugashi", "O'rta", "Zaxira provayder", "Xizmat uzilishini kamaytiradi"],
    ["Parolni oshkor qilish", "Past", "2FA va maslahat", "Hisob buzilishi ehtimoli kamayadi"],
    ["Elektr uzilishi", "O'rta", "UPS va backup", "Ma'lumotni yo'qotish kamayadi"],
    ["Foydalanuvchi xatosi", "Yuqori", "Validator va izoh", "Noto'g'ri amallar kamayadi"],
    ["Prompt injection", "O'rta", "Kontekstni tozalash", "AI xavfsizligi oshadi"],
    ["Zaxira nusxasi buzilishi", "Past", "Tiklash testi", "Tiklash ishonchliligi oshadi"],
]


def append_section(doc, title, paras, *, add_break=True):
    if add_break:
        doc.add_page_break()
    add_heading(doc, title, level=1, align="center")
    for para in paras:
        add_para(doc, para)


def append_subsection(doc, title, paras, *, before_break=False):
    if before_break:
        doc.add_page_break()
    add_para(doc, title, bold=True, align="left", first_line=False, after=4, before=8)
    for para in paras:
        add_para(doc, para)


def build():
    source = pick_source()
    doc = Document(str(source))

    # extend Chapter IV
    doc.add_page_break()
    add_heading(doc, "IV BOB. HAYOT FAOLIYATI XAVFSIZLIGI", level=1, align="center")
    for para in CH4_INTRO:
        add_para(doc, para)

    append_subsection(doc, "4.1. Kompyuter xavfsizligi va sanitariya-gigiyena qoidalari", CH4_1_PARAS[1:])
    add_caption(doc, "4.1.1-jadval. Kompyuter bilan ishlashda uchrashi mumkin bo'lgan xavflar va profilaktika choralari")
    add_table(doc, CH4_1_TABLE[0], CH4_1_TABLE[1:])

    append_subsection(doc, "4.2. Texnik obyektlarda mehnatni muhofaza qilishni tashkil etish", CH4_2_PARAS[1:])
    add_caption(doc, "4.2.1-jadval. Texnik obyektlarda mehnat muhofazasining asosiy yo'nalishlari")
    add_table(doc, CH4_2_TABLE[0], CH4_2_TABLE[1:])

    add_para(doc, "Axborot xavfsizligi va AI xizmatlaridan mas'uliyatli foydalanish", bold=True, align="left", first_line=False, after=4, before=8)
    for para in CH4_3_PARAS:
        add_para(doc, para)
    for para in CH4_SUMMARY:
        add_para(doc, para)

    # conclusion
    append_section(doc, "XULOSA", XULOSA_PARAS)

    # references
    doc.add_page_break()
    add_heading(doc, "FOYDALANILGAN ADABIYOTLAR", level=1, align="center")
    add_para(doc, "Quyida diplom loyihasida foydalanilgan asosiy normativ hujjatlar, ilmiy manbalar va texnik adabiyotlar ro'yxati keltiriladi.", first_line=False)
    for ref in REFERENCES:
        add_para(doc, ref, first_line=True, after=2)

    # appendices
    doc.add_page_break()
    add_heading(doc, "ILOVALAR", level=1, align="center")

    for para in APPENDIX_A:
        add_para(doc, para)

    add_para(doc, "Ilova A uchun texnik buyruqlar namunasi", bold=True, align="left", first_line=False, after=4, before=8)
    add_code_block(doc, [
        "Set-Location \"d:\\SkillMap\\front\"",
        "python -m venv .venv",
        ".\\.venv\\Scripts\\Activate.ps1",
        "pip install -r requirements.txt",
        "python manage.py migrate",
        "python manage.py createsuperuser",
        "python manage.py runserver",
        "npm install",
        "npm run dev",
        "",
        "# production build",
        "npm run build",
        "python manage.py collectstatic --noinput",
    ])

    doc.add_page_break()
    for para in APPENDIX_B:
        add_para(doc, para)
    add_caption(doc, "Ilova B.1-jadval. Asosiy REST API endpointlari katalogi")
    add_table(doc, ["HTTP", "Marshrut", "Vazifa", "Auth"], API_ROWS)

    doc.add_page_break()
    for para in APPENDIX_C:
        add_para(doc, para)
    add_caption(doc, "Ilova C.1-jadval. Onboarding savollari va baholash mezonlari")
    add_table(doc, ["T.R.", "Savol yo'nalishi", "Maqsad", "Baholash"], ONBOARDING_ROWS)

    doc.add_page_break()
    for para in APPENDIX_D:
        add_para(doc, para)
    add_caption(doc, "Ilova D.1-jadval. Xavf va sinovlar uchun tekshiruv ro'yxati")
    add_table(doc, ["T.R.", "Sinov ssenariysi", "Kutilgan natija", "Holat"], TEST_ROWS)
    add_caption(doc, "Ilova D.2-jadval. Xavflar reyestri va javob choralariga qisqa misol")
    add_table(doc, ["Xavf", "Ehtimol", "Chora", "Izoh"], RISK_ROWS)

    doc.add_page_break()
    add_para(doc, f"Ilova D.3. SkillMap ma'lumot oqimining qisqa izohi: foydalanuvchi profili → onboarding → tekshirish → AI router → validator → saqlash → progress qayd", bold=True, align="left", first_line=False, after=4)
    add_code_block(doc, [
        "User profile -> Onboarding -> Validation -> AI Router -> Validator -> Storage",
        "         |             |             |            |          |",
        "         v             v             v            v          v",
        "      JWT auth     Serializer    Retry/Cache   JSON schema  PostgreSQL",
        "",
        "# Har bir qadam audit logda qayd etiladi",
        "# Xatolik yuz bersa fallback generator ishga tushadi",
    ])

    doc.save(str(OUTPUT))
    print(f"OK: {OUTPUT}")


if __name__ == "__main__":
    build()

